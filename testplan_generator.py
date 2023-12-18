#----------------------------------------------------
# Author: Lakshmanan Kuppan
# https://www.linkedin.com/in/lakshmanankuppan/
#----------------------------------------------------

import re
import sys
import markdown
from palm_api import send_query_to_ai
from palm_api import calculate_token_size
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import Pt

def set_prompt(title, text):
    prompt = '''
Given the functional specification details provided below, generate a test plan in Markdown format incrementally as you receive each chunk. Each chunk should include a part of the test plan with the following structure:

### Testcase ID: [integer]
#### Testcase Title
- [string]
#### Test Topology
- [string]
#### Test Procedure
1. [string]
2. [string]
3. [string]
#### Pass Criteria
- [string]

{}
{} '''.format(title, text)
    return prompt

def convert_docx_to_text(FS_docx):
    try:
        doc = Document(FS_docx)
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        return "\n".join(content)
    except:
        print("ERROR in word document")
        sys.exit()

def get_testplan_chunks(testplan_docx):
    text = convert_docx_to_text("FS.docx")
    text_len = calculate_token_size(text)
    token_limit = 3000  # Adjust based on the model's limit
    text_chunks = [text[i:i+token_limit] for i in range(0, text_len, token_limit)]
    return text_chunks

def generate_testplan(text_chunks):

    cnt = 0
    test_plan = []
    title = "EXOS Feature Functional Specification"

    print("Sending chunk by chunk ", end='', flush=True)
    for chunk in text_chunks :
        cnt +=1
        print(f'..{cnt}', end='', flush=True)
        prompt = set_prompt(title, chunk)
        markdown_testplan = send_query_to_ai(prompt)
        test_plan.append(markdown_testplan)
    return test_plan

def update_testcase_id(test_plan):
    updated_testplan=[]
    tc_section=1
    for each_tc in test_plan:
        #tcid = re.search(r'Testcase ID: (\d+)', each_tc)[1]
        #new_tcid = f'{tc_section}.{tcid}'
        testcase = re.sub(r'### Testcase ID: (\d+)', fr'\n### Testcase ID: {tc_section}.\1', each_tc)
        updated_testplan.append(testcase)
        tc_section+=1
    return updated_testplan

def convert_markdown_2_html(html_file, testplan):

    try:
        htmloutput = [markdown.markdown(tp) for tp in testplan]
        html_testplan = "\n".join(htmloutput)

        with open(html_file, "w", encoding="utf-8") as fid:
            fid.write(html_testplan)
        return True
    except:
        return False


def create_word_from_html(html, doc):
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol']):
        if element.name == 'h1':
            paragraph = doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            paragraph = doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            paragraph = doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            paragraph = doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                paragraph = doc.add_paragraph('\u2022 ' + li.get_text(), style='ListBullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                paragraph = doc.add_paragraph('1. ' + li.get_text(), style='ListNumber')

        # Adjust font size if needed
        for run in paragraph.runs:
            run.font.size = Pt(12)

def create_testplan(functional_spec, html_file):
    text_chunks = get_testplan_chunks(functional_spec)
    test_plan = generate_testplan(text_chunks)
    updated_testplan = update_testcase_id(test_plan)
    if convert_markdown_2_html(html_file, updated_testplan) :
        print(f'\n Testplan File created in html format {html_file}')
        return True
    else:
        print(f'Markdown to HTML Failed')
        return False

if __name__ == '__main__':
    functional_spec = "FS.docx"
    testplan_in_html_format = "testplan.html"
    create_testplan(functional_spec, testplan_in_html_format)
